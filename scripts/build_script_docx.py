# -*- coding: utf-8 -*-
"""Generate Script_Soutenance.docx — 16 slides, ~17 minutes, simple English."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY  = RGBColor(0x00, 0x22, 0x55)
GOLD  = RGBColor(0xC8, 0xA0, 0x32)
GREY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1A, 0x73, 0x4A)

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────
def set_run(run, bold=False, italic=False, color=None, size=None):
    run.bold   = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run(run, bold=True, color=NAVY, size=20)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '002255')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def slide_heading(number, title, speaker, timing):
    """Slide section header: number + title in navy, speaker chip in gold, timing in grey."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"SLIDE {number} — {title.upper()}   ")
    set_run(r1, bold=True, color=NAVY, size=13)
    r2 = p.add_run(f"[{speaker}]")
    set_run(r2, bold=True, color=GOLD, size=12)
    r3 = p.add_run(f"  · {timing}")
    set_run(r3, italic=True, color=GREY, size=11)
    return p

def body(text, italic=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    set_run(run, italic=italic, color=color or RGBColor(0x21, 0x25, 0x29))
    return p

def transition(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.4)
    r = p.add_run(f"→  {text}")
    set_run(r, italic=True, color=GREY, size=10)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        # navy background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '002255')
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = 'EEF3F9' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            p = cell.paragraphs[0]
            p.add_run(cell_text).font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════
# DOCUMENT HEADER
# ══════════════════════════════════════════════════════════════
heading1("ORAL DEFENSE SCRIPT — WORD FOR WORD")
p = doc.add_paragraph()
r = p.add_run("The Role of Weather in French Day-Ahead Electricity Price Forecasting")
set_run(r, bold=True, color=NAVY, size=12)
p = doc.add_paragraph()
r = p.add_run("Leo Cambreleng & Lyam Oumedjeber  ·  EDHEC MSc DAAI  ·  June 2026")
set_run(r, color=GREY, size=10)
p = doc.add_paragraph()
r = p.add_run("FORMAT: 20 min presentation + Q&A  ·  LEO speaks slides 1–7  ·  LYAM speaks slides 8–15  ·  Slide 16: both  ·  Target: ~17 min")
set_run(r, italic=True, color=GREY, size=10)
divider()

# ══════════════════════════════════════════════════════════════
# SLIDES
# ══════════════════════════════════════════════════════════════

slide_heading(1, "Title", "LEO", "~20 s")
body("Good morning, Professor. I'm Leo Cambreleng, and with my co-author Lyam Oumedjeber, we're presenting our Master's thesis: \"The Role of Weather in French Day-Ahead Electricity Price Forecasting.\" I'll cover the first half, then hand over to Lyam.")
divider()

slide_heading(2, "Context & Motivation", "LEO", "~1 min 30 s")
body("Why forecast French electricity prices — and why might weather help?")
body("First: electricity can't be stored. Supply and demand balance every hour, creating extreme volatility — prices can spike to several thousand euros per megawatt-hour.")
body("Second: France is unique. Around 70% of output is nuclear, and because heating is mostly electric, a one-degree temperature drop adds about 2.4 gigawatts of demand — the highest thermosensitivity in Europe.")
body("Third: the stakes are real. A one-euro improvement on a 100-megawatt book is worth about 876,000 euros per year.")
body("The question is: does weather actually help — or is that signal already captured by other variables?")
divider()

slide_heading(3, "Research Question", "LEO", "~1 min")
body("We structured the thesis around four sub-questions.")
body("Can machine learning beat a naïve statistical benchmark? Do weather features add significant value beyond fundamentals alone? Does that value change between a stable market and a crisis — and why? And do better forecasts translate into real trading profit?")
body("These four questions map directly onto the slides ahead.")
divider()

slide_heading(4, "Data", "LEO", "~1 min")
body("Our dataset covers January 2018 to April 2025 — around 64,000 hourly observations from four public sources.")
body("ENTSO-E for prices, load forecast, generation, and cross-border flows. ERA5 from ECMWF for weather — temperature, wind, solar, precipitation. TTF gas and ARA coal for fuel prices. In total, 35 engineered features.")
body("One key point: the data spans three very different regimes — pre-crisis, the 2022 energy crisis where prices briefly exceeded 1,000 euros, and post-crisis normalisation. That heterogeneity is central to our results.")
divider()

slide_heading(5, "Methodology", "LEO", "~1 min")
body("We used a four-model ablation design.")
body("Model A is the naïve benchmark: the price from the same hour last week. Model B is a Random Forest on 27 features — no weather. Model C is our main model: Random Forest with all 35 features including weather. Model D is XGBoost with the same 35 features.")
body("The key comparison is C versus B: same architecture, same data — only the weather features differ. This isolates weather's marginal contribution, tested out-of-sample with the Diebold-Mariano test.")
divider()

slide_heading(6, "Results — Stable Regime", "LEO", "~1 min 30 s")
body("On the stable test period — May 2024 to April 2025, 8,640 observations — the naïve benchmark gives a MAE of 33 euros and an R-squared of 0.16.")
body("Both Random Forest models cut that error by 49%, bringing MAE to 16.9 euros and R-squared to 0.78.")
body("Now the critical number: Model C with weather, 16.94. Model B without weather, 16.95. One hundredth of a euro. XGBoost underperforms at 19.14.")
body("Machine learning clearly works — but weather, in this regime, adds essentially nothing.")
divider()

slide_heading(7, "Central Finding", "LEO", "~1 min")
body("The statistical test confirms it. The Diebold-Mariano test for C versus B gives p = 0.572 — not significant.")
body("But test either RF model against the naïve benchmark, and the DM statistic is around minus 53, with p below 0.001. Machine learning adds massive, unambiguous value. Weather does not — in this regime.")
body("Why? And is it always the case? I hand over to Lyam.")
transition("\"I'll hand over to Lyam, who will explain the mechanism and then show what happens in the 2022 crisis.\"")
divider()

slide_heading(8, "Feature Importance", "LYAM", "~1 min")
body("Thank you, Leo. The feature importance chart explains why weather is redundant.")
body("Price lags dominate: the 24-hour lag alone contributes 21% of importance, and lags plus rolling means account for about 70%. Then fuel prices: TTF gas 8%, ARA coal 6%, nuclear availability 2%. All weather variables combined: just 2.2%.")
body("Weather's signal is already captured by other features — the next slide explains exactly how.")
divider()

slide_heading(9, "Information Redundancy Hypothesis", "LYAM", "~1 min 30 s")
body("We call this the Information Redundancy Hypothesis. Three channels explain it.")
body("One: the RTE load forecast is built on temperature data. Including it means the model is already conditioning on weather-driven demand. Raw temperature adds nothing new.")
body("Two: TTF gas prices aggregate pan-European heating demand. When it's cold across Europe, gas demand rises, TTF rises. The commodity market does the weather encoding for us.")
body("Three: in France, nuclear availability explains so much price variance that little remains for weather to explain.")
body("These three channels make weather redundant — as long as they're intact. In 2022, they were not.")
transition("\"But this only holds when those three channels are working. In 2022, they all broke at once.\"")
divider()

slide_heading(10, "2022 Crisis Reversal", "LYAM", "~1 min 30 s")
body("We retrained on 2018–2021 and tested on the full year 2022 — a strict temporal split.")
body("The contrast is stark. Stable period: DM statistic minus 0.565, p = 0.572 — not significant. Crisis 2022: DM statistic minus 13.27, p below 0.001.")
body("All models degrade sharply: naïve MAE goes from 33 to 73 euros, RF from 17 to 67. But Model C now beats Model B by 0.73 euros per megawatt-hour — statistically huge because it's consistent across nearly 8,760 hours.")
body("XGBoost degrades the most at 76 euros MAE — confirming RF as the more robust architecture under stress.")
divider()

slide_heading(11, "Why the Crisis Breaks Redundancy", "LYAM", "~1 min 30 s")
body("All three redundancy channels break simultaneously in 2022.")
body("One: TTF decouples from local weather. The Russian gas shock drove prices by geopolitics, not temperature. TTF no longer encodes the weather signal.")
body("Two: nuclear constraints amplify the temperature effect. With availability at just 40% — around 25 of 63 gigawatts — there was no buffer. Every cold degree fed directly into price.")
body("Three: signal-to-noise rises. When prices swing by hundreds of euros, the direct temperature effect becomes large enough to measure statistically.")
body("Takeaway: weather's value is not fixed — it's regime-dependent.")
divider()

slide_heading(12, "Economic Value — Trading Backtest", "LYAM", "~1 min")
body("Do better forecasts make money? We ran a day-ahead directional strategy at 0.30 euros per megawatt-hour — the central cost scenario.")
body("The naïve strategy earns 85,000 euros per megawatt but with a drawdown of 2,367 euros and a Calmar ratio of 37. Both RF models earn around 137,000, with a drawdown five times smaller — 422 euros — and a Calmar of 328.")
body("XGBoost earns slightly more in gross profit but its drawdown is 1,351 euros, three times RF's. RF posts positive Sharpe ratios every single month. Real skill, not luck.")
divider()

slide_heading(13, "Robustness Checks", "LYAM", "~30 s")
body("Quick robustness check. RF's Sharpe ratio across the three cost scenarios moves from 19.51 to 19.16 — under two percent degradation. Positive Sharpe every month. The DM non-significance result holds across multiple sub-periods. The conclusions don't depend on our specific cost assumptions.")
divider()

slide_heading(14, "Limitations & Future Work", "LYAM", "~45 s")
body("Six honest limitations.")
body("ERA5 is perfect hindsight weather — in production you'd use forecasts with 10–30% error, so our result is an upper bound. We didn't do walk-forward retraining. Carbon prices were excluded. XGBoost wasn't tuned — Bayesian optimisation could close the gap. We only cover France; other markets may differ. And transaction costs are flat — large books would face liquidity constraints.")
body("These define the perimeter of what we can claim.")
divider()

slide_heading(15, "Conclusions", "LYAM", "~1 min 30 s")
body("Four takeaways.")
body("One: machine learning decisively beats the naïve benchmark — minus 49% MAE, R-squared 0.78, confirmed across all twelve months.")
body("Two: the value of weather is regime-dependent. Non-significant in stable conditions, highly significant in the 2022 crisis. That's our core contribution.")
body("Three: Random Forest outperforms XGBoost on risk-adjusted metrics — Calmar 328 versus 107 — and holds up better under stress.")
body("Four: the forecasts create real economic value — 137,000 euros per megawatt per year, drawdown five times smaller than the benchmark.")
body("The practical implication: don't always include or always exclude weather. Build a regime detector. Thank you very much.")
divider()

slide_heading(16, "Questions", "LEO + LYAM", "~20 s")
body("[Smile. Pause 2–3 seconds. Let the examiner speak first.]", italic=True, color=GREY)
body("[If silence: \"Would you like us to start with any particular aspect of the methodology?\"]", italic=True, color=GREY)
divider()

# ══════════════════════════════════════════════════════════════
# TIMING TABLE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading1("TIMING REFERENCE")
add_table(
    headers=["Slide", "Speaker", "Target"],
    rows=[
        ("1 — Title",                "LEO",       "20 s"),
        ("2 — Context & Motivation", "LEO",       "1 min 30 s"),
        ("3 — Research Question",    "LEO",       "1 min"),
        ("4 — Data",                 "LEO",       "1 min"),
        ("5 — Methodology",          "LEO",       "1 min"),
        ("6 — Results stable",       "LEO",       "1 min 30 s"),
        ("7 — Central Finding",      "LEO",       "1 min"),
        ("8 — Feature Importance",   "LYAM",      "1 min"),
        ("9 — Redundancy",           "LYAM",      "1 min 30 s"),
        ("10 — 2022 Crisis",         "LYAM",      "1 min 30 s"),
        ("11 — Why 2022",            "LYAM",      "1 min 30 s"),
        ("12 — Trading Backtest",    "LYAM",      "1 min"),
        ("13 — Robustness",          "LYAM",      "30 s"),
        ("14 — Limitations",         "LYAM",      "45 s"),
        ("15 — Conclusions",         "LYAM",      "1 min 30 s"),
        ("16 — Questions",           "LEO + LYAM","20 s"),
        ("TOTAL",                    "",          "~17 min"),
    ],
    col_widths_cm=[8.0, 3.5, 3.0]
)

# ══════════════════════════════════════════════════════════════
# KEY TRANSITIONS
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading1("KEY TRANSITIONS")

p = doc.add_paragraph()
r = p.add_run("Slide 7 → 8  (LEO hands to LYAM):")
set_run(r, bold=True, color=NAVY, size=11)
body('"I\'ll hand over to Lyam, who will explain the mechanism and show what happens in the 2022 crisis."')

p = doc.add_paragraph()
r = p.add_run("Slide 9 → 10  (LYAM continues):")
set_run(r, bold=True, color=NAVY, size=11)
body('"But this only holds when those three channels are working. In 2022, they all broke at once."')

p = doc.add_paragraph()
r = p.add_run("End of slide 15:")
set_run(r, bold=True, color=NAVY, size=11)
body('"Thank you very much. We\'re happy to take your questions."')

# ── Save ──────────────────────────────────────────────────────
OUT = r"C:\Users\Public\fr-power-thesis\outputs\Script_Soutenance.docx"
doc.save(OUT)
print("Saved:", OUT)
