# -*- coding: utf-8 -*-
"""
Guide de comprehension complet de la these (pour Leo & Lyam).
Plain, tres structure, beaucoup de bullets. FR avec termes EN entre parentheses.
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0, 50, 100)
DGREY = RGBColor(60, 60, 60)
BLACK = RGBColor(0, 0, 0)

doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(4)


def _runs(p, text):
    """Parse **bold** markup into runs."""
    for i, seg in enumerate(re.split(r"(\*\*.*?\*\*)", text)):
        if not seg:
            continue
        r = p.add_run(seg[2:-2] if seg.startswith("**") else seg)
        if seg.startswith("**"):
            r.bold = True


def H1(text):
    p = doc.add_paragraph(); p.space_before = Pt(16); p.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8'); bot.set(qn('w:space'), '3'); bot.set(qn('w:color'), '003264')
    pb.append(bot); pPr.append(pb)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY


def H2(text):
    p = doc.add_paragraph(); p.space_before = Pt(11); p.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY


def H3(text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = DGREY


def P(text, after=5):
    p = doc.add_paragraph(); p.space_after = Pt(after)
    _runs(p, text); return p


def B(text, lvl=0):
    style = "List Bullet" if lvl == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style); p.space_after = Pt(2)
    _runs(p, text); return p


def NB(text, n):
    p = doc.add_paragraph(style="List Number"); p.space_after = Pt(2)
    _runs(p, text); return p


def callout(label, text):
    p = doc.add_paragraph(); p.space_before = Pt(4); p.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), 'EEF2F7'); pPr.append(sh)
    bd = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement(f'w:{side}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '4'); e.set(qn('w:color'), '003264'); bd.append(e)
    pPr.append(bd)
    r = p.add_run(label + " "); r.bold = True; r.font.color.rgb = NAVY
    _runs(p, text)


# ===================== COVER =====================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("GUIDE DE COMPRÉHENSION COMPLET"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Tout comprendre de A à Z pour la soutenance"); r.italic = True; r.font.size = Pt(13); r.font.color.rgb = DGREY
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; t.space_after = Pt(2)
r = t.add_run("« The Role of Weather in French Day-Ahead Electricity Price Forecasting »"); r.font.size = Pt(11)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Leo Cambreleng & Lyam Oumedjeber — EDHEC MSc Data Analysis & AI"); r.font.size = Pt(10); r.font.color.rgb = DGREY
P("", 2)
callout("Comment lire ce guide :",
        "il est fait pour être lu **dans l'ordre** une fois en entier (compréhension), puis relu par section avant la "
        "soutenance. Les termes techniques sont donnés en français avec **l'équivalent anglais entre parenthèses** — "
        "c'est exactement le mot que tu prononceras à l'oral. Tout est en bullets pour être révisable vite. "
        "Les chiffres clés à connaître par cœur sont à la toute fin (Partie X).")

# ===================== 0. PITCH =====================
H1("0.  LE SUJET EN 1 MINUTE (le pitch à avoir en tête)")
P("Si quelqu'un te réveille à 3h du matin et te demande « c'est quoi ta thèse ? », voici la réponse :")
callout("En une phrase :",
        "On essaie de **prédire le prix de l'électricité en France pour le lendemain**, avec des modèles de "
        "machine learning, et on se demande surtout : **est-ce que la météo aide à mieux prédire ?** Réponse : "
        "**ça dépend de la situation du marché**. En période normale, non (l'info météo est déjà cachée ailleurs). "
        "En période de crise (2022), oui, clairement.")
P("Les 4 messages que tu dois pouvoir défendre :")
NB("**Le machine learning marche** : nos modèles divisent l'erreur par ~2 par rapport à une méthode bête (naïve).", 1)
NB("**La météo est \"regime-dependent\"** : inutile en régime calme, utile en crise. → **C'est notre vraie trouvaille.**", 2)
NB("**Random Forest > XGBoost** : plus robuste, moins risqué.", 3)
NB("**Ça rapporte de l'argent** : une stratégie de trading basée sur nos prévisions est rentable.", 4)

# ===================== PARTIE I =====================
H1("PARTIE I — LE CONTEXTE : LE MARCHÉ DE L'ÉLECTRICITÉ")

H2("1.  Pourquoi l'électricité n'est pas un produit comme les autres")
B("**On ne peut pas la stocker** (à grande échelle, économiquement). Le pétrole, le blé, l'or : on stocke. "
  "L'électricité : non. Conséquence énorme → **à chaque seconde, la production doit exactement égaler la consommation.**")
B("Comme on ne peut pas faire de stock pour lisser, **le prix doit ajuster l'équilibre en temps réel** → ça crée "
  "une **volatilité extrême** (le prix peut être négatif le matin et à 1000 €/MWh le soir).")
B("Unité de prix : **€/MWh** (euros par mégawattheure). 1 MWh = consommation d'environ 1000 foyers pendant 1h "
  "(ordre de grandeur). Prix « normal » en France ≈ 50–100 €/MWh.")
B("**Saisonnalités fortes et régulières** : par heure (pic le matin et le soir), par jour (semaine ≠ week-end), "
  "par saison (hiver = chauffage = demande haute). C'est cette régularité qui rend la prévision possible.")

H2("2.  Comment se forme un prix de l'électricité (le « merit order »)")
B("On empile les moyens de production **du moins cher au plus cher** (ça s'appelle le **merit order** / ordre de mérite).")
B("Ordre typique du moins cher au plus cher : **renouvelables (vent, solaire) → nucléaire → charbon → gaz**. "
  "Les renouvelables ont un coût marginal ~0 (le vent est gratuit), le gaz est cher (il faut acheter le gaz).")
B("On sert la demande en remontant cette pile. **La dernière centrale allumée pour satisfaire le dernier MWh "
  "fixe le prix pour TOUT LE MONDE** : c'est la centrale « marginale » (marginal plant).")
B("**En France, hors heures de pointe, le nucléaire couvre la base.** Mais quand la demande monte (hiver, soir), "
  "c'est souvent une **centrale à gaz** qui est marginale → **le prix de l'élec suit alors le prix du gaz.** "
  "→ Retiens ça, c'est central pour comprendre le rôle du gaz (TTF) plus loin.")

H2("3.  Le marché « day-ahead » EPEX SPOT (notre terrain de jeu)")
B("**Day-Ahead (DA) = marché \"pour le lendemain\".** Aujourd'hui (jour D) on achète/vend l'électricité de "
  "demain (jour D+1), heure par heure (24 prix par jour).")
B("Géré par **EPEX SPOT** (la bourse de l'élec pour la France/Allemagne...). C'est **notre cible de prévision** : "
  "le prix de clôture du marché day-ahead France.")
B("**Mécanisme = enchère uniforme (uniform-price auction)** : tout le monde envoie ses offres d'achat et de vente "
  "AVANT **12h00 CET** (la « gate closure » / fermeture du guichet). À 12h, l'opérateur calcule **un seul prix "
  "d'équilibre par heure**, et tous ceux qui ont été retenus échangent à ce prix unique.")
callout("Implication clé pour le trading (slide backtest) :",
        "le signal pour chaque heure de demain doit être calculé **uniquement avec l'info disponible avant 12h "
        "aujourd'hui**. Notamment : on connaît déjà les 24 prix d'aujourd'hui. C'est ce qui rend notre stratégie "
        "« exécutable » (executable) et pas de la triche rétrospective.")

H2("4.  Les spécificités françaises (à marteler — c'est l'angle du sujet)")
B("**~63 GW de nucléaire installé**, soit **~70 % de la production** annuelle. → La **disponibilité du parc "
  "nucléaire** (nuclear availability) est LE déterminant n°1 des prix français.")
B("**Chauffage très électrique** : beaucoup de foyers se chauffent à l'élec. → **Thermosensibilité ≈ 2,4 GW par °C** "
  "en hiver (la demande monte de 2,4 GW quand il fait 1°C de moins). **C'est la plus forte d'Europe.** "
  "→ C'est exactement pour ça que la météo POURRAIT compter en France.")
B("Conséquence : en France, le prix dépend d'un duel **demande (pilotée par la température) vs offre nucléaire "
  "(disponible ou pas)**.")

H2("5.  Pourquoi les prix sont si « bizarres » (volatilité, pics, prix négatifs)")
B("**Pics de prix (spikes)** : panne soudaine d'une centrale + vague de froid → prix qui explose. Ces événements "
  "sont **imprévisibles par nature** → même un bon modèle a du mal (on en reparle dans les limites).")
B("**Prix négatifs** : oui, ça existe. Quand des producteurs « inflexibles » (nucléaire, éolien) sont forcés de "
  "produire et qu'on ne peut pas les arrêter facilement, et que la demande est faible → ils **paient pour qu'on "
  "consomme** leur surplus. ~3 à 8 % des heures récentes.")
B("Problème technique que ça crée : certaines métriques d'erreur (comme le **MAPE**) explosent quand le prix est "
  "proche de 0 ou négatif → on les évite (voir Partie V).")

# ===================== PARTIE II =====================
H1("PARTIE II — QU'EST-CE QUI FAIT BOUGER LES PRIX")
P("Pour prédire un prix, il faut nourrir le modèle avec les bonnes « variables explicatives » (features). Voici "
  "les grandes familles de ce qui influence le prix.")

H2("6.  Les fondamentaux (fundamentals)")
B("**La charge / demande (load)** : combien on consomme. Surtout : la **prévision de charge (load forecast)** "
  "publiée par RTE la veille. **C'est une de nos features les plus importantes** (et un point CENTRAL du sujet, "
  "voir redondance plus loin).")
B("**La production par source (generation by source)** : nucléaire, gaz, hydraulique (fil de l'eau + réservoir), "
  "solaire, éolien. Plus il y a de renouvelable/nucléaire dispo, plus le prix tend à baisser.")

H2("7.  Les combustibles (fuel prices) — surtout le gaz")
B("**TTF (Title Transfer Facility)** = le prix de référence du **gaz naturel** en Europe (€/MWh). **C'est notre "
  "feature de combustible la plus importante.** Pourquoi ? Parce que la centrale à gaz est souvent la centrale "
  "marginale → **prix élec ≈ suit le prix du gaz** aux heures de pointe.")
B("**Charbon ARA (ARA coal)** : prix du charbon livré à Amsterdam-Rotterdam-Anvers (€/tonne). Concurrent du gaz "
  "pour la production thermique.")
B("**EUA / quotas carbone (EU ETS)** : on ne les a PAS pris comme variable séparée. Justification (à connaître, "
  "Milos peut demander) : leur signal est **déjà largement contenu dans l'écart gaz–charbon** qu'on a. "
  "**On l'assume comme une limite.**")

H2("8.  Les flux transfrontaliers (cross-border flows)")
B("La France est connectée à l'**Allemagne, l'Espagne, le Royaume-Uni, la Belgique**. L'élec circule selon les "
  "écarts de prix (on importe quand c'est moins cher ailleurs, on exporte sinon).")
B("On utilise les **flux nets** (net flow) avec ces 4 voisins comme features.")

H2("9.  La météo (weather) — le cœur du sujet")
B("**Température (temperature_2m)** : froid → chauffage → demande ↑ → prix ↑. Chaud → clim (l'été) → demande ↑ aussi.")
B("**Vent (wind speed)** : vent ↑ → production éolienne ↑ → offre ↑ → prix ↓.")
B("**Rayonnement solaire (solar radiation)** : soleil ↑ → production solaire ↑ → prix ↓ (surtout midi).")
B("**Précipitations (precipitation)** : influence l'hydraulique.")
B("Donc la météo agit sur **les deux côtés** : sur la **demande** (température) ET sur **l'offre renouvelable** "
  "(vent, soleil). D'où l'intuition de départ : la météo devrait aider à prédire.")

H2("10.  HDD et CDD (les « degrés-jours », à savoir expliquer)")
B("**HDD (Heating Degree Days, degrés-jours de chauffage)** : mesure du besoin de **chauffage**. Si la température "
  "passe sous un **seuil** (ici **17 °C**, le seuil RTE), on compte l'écart. Plus il fait froid, plus HDD est grand "
  "→ plus on chauffe → demande ↑.")
B("**CDD (Cooling Degree Days, degrés-jours de climatisation)** : symétrique pour la **clim**, au-dessus d'un seuil "
  "(ici **22 °C**). Plus il fait chaud, plus CDD est grand → clim ↑.")
B("Pourquoi un seuil ? Parce que la relation température→demande **n'est pas linéaire** : entre 17 et 22 °C, on "
  "ne chauffe ni ne climatise → effet ~nul. HDD/CDD capturent cette **non-linéarité à seuil**.")
B("**WSI (Weather Stress Index, indice de stress météo)** : variable qu'on a fabriquée = **HDD / vitesse du vent** "
  "= capture le scénario le pire (**froid + pas de vent** = forte demande + peu d'éolien).")

# ===================== PARTIE III =====================
H1("PARTIE III — LA DATA (les données)")

H2("11.  Les 4 sources de données")
B("**ENTSO-E** (European Network of Transmission System Operators for Electricity) : plateforme publique "
  "(Transparency Platform). On en tire : **prix day-ahead** (la cible), **prévision de charge**, **production par "
  "source**, **flux transfrontaliers**. Récupéré via la lib Python **entsoe-py**.")
B("**ERA5** : données météo de **réanalyse (reanalysis)** produites par l'ECMWF (Copernicus). Grille de 0,25°. "
  "On prend une boîte couvrant la France métropolitaine et on fait une **moyenne spatiale**. Récupéré via **cdsapi**.")
B("**TTF gaz** et **charbon ARA** : prix des combustibles, sources de marché publiques.")
B("Période : **janvier 2018 → avril 2025**, au pas **horaire (hourly)** → **~64 000 lignes** (observations).")
callout("Point de vigilance (Milos peut demander) :",
        "ERA5 est de la **réanalyse** = la météo **réellement observée** a posteriori, **pas une prévision météo**. "
        "Donc on teste la valeur d'une **météo parfaite** (best case). Si même une météo parfaite est inutile en "
        "régime calme, alors une vraie prévision (imparfaite) le sera encore plus → **ça renforce notre conclusion** "
        "en régime stable. En crise, c'est une nuance à reconnaître honnêtement.")

H2("12.  C'est quoi une « série temporelle » et pourquoi c'est particulier")
B("**Série temporelle (time series)** = des données ordonnées dans le temps (1 valeur par heure ici). "
  "L'ordre compte : la valeur d'hier aide à prédire celle de demain (**autocorrélation**).")
B("Conséquence pour le ML : on **ne peut pas mélanger** passé et futur n'importe comment. → split temporel strict "
  "(Partie IV).")

H2("13.  Le « feature engineering » (fabriquer les variables)")
P("À partir des données brutes, on construit **35 variables (features)** en 5 familles. Les techniques importantes :")
H3("a) Les retards / lags (les plus importants !)")
B("**Lag = la valeur du prix il y a X heures.** Ex : **price_lag24h** = le prix à la même heure hier. "
  "**price_lag168h** = même heure il y a 168h = **il y a 1 semaine** (168 = 24×7).")
B("Pourquoi c'est puissant : le prix de l'élec est très **autocorrélé**. Le meilleur indice du prix de demain "
  "8h, c'est souvent le prix d'aujourd'hui 8h (lag-24h) ou de la semaine dernière 8h (lag-168h).")
H3("b) Les moyennes glissantes / rolling means")
B("**rolling mean** = moyenne des X dernières heures (décalée pour ne pas tricher). Lisse le bruit et donne "
  "le « niveau » récent du prix.")
H3("c) L'encodage cyclique (sin/cos)")
B("Problème : l'heure 23 et l'heure 0 sont **voisines** dans la réalité, mais 23 et 0 sont « loin » numériquement. "
  "Pareil décembre (12) et janvier (1).")
B("Solution : on encode l'heure et le mois avec **sinus et cosinus** → ça transforme le temps en cercle, donc "
  "23h et 0h redeviennent proches. C'est le **cyclic encoding**.")
H3("d) Les variables dérivées (derived)")
B("**wind_power_proxy** = vitesse du vent au cube (la puissance éolienne ∝ vent³).")
B("**WSI** = HDD / vent (vu plus haut). **nuclear_avail_ratio** = production nucléaire / 63 000 MW installés.")

# ===================== PARTIE IV =====================
H1("PARTIE IV — LES MODÈLES (le cœur machine learning)")

H2("14.  C'est quoi « prédire » en machine learning")
B("**Apprentissage supervisé (supervised learning)** : on montre au modèle plein d'exemples passés où on connaît "
  "à la fois les **entrées X** (les 35 features) et la **bonne réponse y** (le prix réalisé). Le modèle apprend "
  "la relation **X → y**.")
B("Ensuite on lui donne des X **qu'il n'a jamais vus** et on regarde si sa prédiction ŷ est proche du vrai y.")
B("**Régression (regression)** : ici on prédit un **nombre** (un prix), pas une catégorie → c'est de la régression.")

H2("15.  Train / Test et pourquoi un split TEMPOREL strict")
B("**Train set (jeu d'entraînement)** : les données passées sur lesquelles le modèle apprend.")
B("**Test set (jeu de test)** : des données **mises de côté**, jamais vues à l'entraînement, pour juger honnêtement.")
B("**Crucial en série temporelle : on entraîne sur le PASSÉ, on teste sur le FUTUR.** On ne mélange jamais. "
  "Sinon **fuite de données (data leakage)** = le modèle « voit » le futur → résultats faussement bons.")
B("Nos 2 jeux de test : **(1)** mai 2024 → avril 2025 (régime **stable**, 8 640 h) ; **(2)** toute l'année **2022** "
  "(la **crise**, 8 760 h), avec entraînement seulement sur 2018–2021.")

H2("16.  Modèle A — le benchmark naïf (la barre à battre)")
B("**Naïf (naïve)** = la prévision la plus bête possible : **« demain sera comme la semaine dernière »**. "
  "Formellement : prix prévu = prix de la même heure il y a 168h (lag-168h).")
B("Pourquoi 168h et pas 24h ? Parce que le motif **hebdomadaire** est plus régulier que le motif journalier "
  "(un mardi ressemble plus au mardi d'avant qu'au lundi d'avant). → benchmark **plus exigeant**.")
B("**À quoi ça sert ?** À donner un **point de référence**. Si nos modèles ML ne battent pas ça, ils sont inutiles.")

H2("17.  Un arbre de décision (decision tree) — la brique de base")
B("Un **arbre** pose une suite de questions oui/non sur les features : « load > 60 GW ? » → oui/non → "
  "« température < 5°C ? » → etc. Au bout (la **feuille / leaf**), il donne une prédiction de prix.")
B("Atout : il capture naturellement les **seuils et non-linéarités** (ex : l'effet du froid sous 17°C). "
  "Défaut : un seul arbre **surapprend (overfit)** = il colle trop au passé et généralise mal.")

H2("18.  Random Forest (RF) — notre modèle principal")
B("**Idée : au lieu d'un arbre, on en fait 500, et on fait la MOYENNE de leurs prédictions.** « Forêt » = "
  "ensemble d'arbres (ensemble method).")
B("Pour que les arbres soient **différents** entre eux (sinon ça ne sert à rien de les moyenner) :")
B("**Bootstrap (bagging)** : chaque arbre est entraîné sur un **tirage aléatoire avec remise** des données.", 1)
B("**Sous-ensemble de features** : à chaque question, l'arbre ne regarde qu'**un sous-ensemble aléatoire** "
  "des variables (règle **sqrt** = racine carrée du nombre de features).", 1)
B("**Résultat : la moyenne de 500 arbres « décorrélés » réduit la variance** (le surapprentissage) sans devenir "
  "trop biaisée. C'est robuste et stable.")
B("Réglages (hyperparamètres) : **500 arbres**, profondeur max **10**, min **5** obs par feuille.")
H3("Pourquoi RF est parfait pour l'élec (3 raisons à citer)")
B("**1.** Capture les **non-linéarités à seuil** (température, etc.).")
B("**2.** **Robuste aux valeurs extrêmes** (les pics de prix ne le déstabilisent pas trop).")
B("**3.** Donne une **mesure d'importance des variables (MDI)** → interprétable, et c'est exactement ce qu'il "
  "nous faut pour répondre « est-ce que la météo compte ? ».")

H2("19.  XGBoost — le concurrent (boosting)")
B("**Boosting** = autre façon de combiner des arbres. Au lieu de les faire en parallèle et moyenner (RF), on les "
  "fait **en séquence** : chaque nouvel arbre **corrige les erreurs** du précédent.")
B("Très puissant en général, mais **plus sensible aux valeurs extrêmes** : il peut « sur-corriger » sur les pics "
  "de prix bizarres. → chez nous il est **moins bon et moins robuste** que RF (surtout en crise 2022).")
B("Réglages : 500 arbres, **learning rate 0,05**, profondeur **6**, subsample 0,8.")

H2("20.  Le design A / B / C / D et l'idée d'ablation")
B("**A** = naïf. **B** = RF **sans météo** (27 features). **C** = RF **avec météo** (35 features, **modèle principal**). "
  "**D** = XGBoost avec météo (35 features).")
B("**Étude d'ablation (ablation study)** = on enlève UNE chose et on regarde l'impact. Ici **C vs B** : seule "
  "différence = la météo. → **comparer C et B isole EXACTEMENT la valeur ajoutée de la météo.** C'est le design "
  "qui rend notre question testable proprement.")

# ===================== PARTIE V =====================
H1("PARTIE V — MESURER LA PERFORMANCE")

H2("21.  Les métriques d'erreur (à savoir expliquer une par une)")
B("**MAE (Mean Absolute Error, erreur absolue moyenne)** : moyenne de |vrai − prévu|. En €/MWh. **Notre métrique "
  "principale** car **robuste aux extrêmes**. « En moyenne on se trompe de X €/MWh ».")
B("**RMSE (Root Mean Square Error)** : racine de la moyenne des erreurs **au carré**. **Pénalise plus les grosses "
  "erreurs** → sensible aux pics. Si RMSE >> MAE, c'est qu'il y a quelques grosses erreurs (les spikes).")
B("**R² (coefficient de détermination)** : part de la variance du prix **expliquée** par le modèle. **0 = nul** "
  "(aussi bien que prédire la moyenne), **1 = parfait**. On passe de 0,16 (naïf) à **0,78** (RF) → on explique 78 %.")
B("**sMAPE (symmetric Mean Absolute Percentage Error)** : erreur en **%**, version « symétrique » qui **évite la "
  "division par zéro** quand le prix est ~0 ou négatif.")
B("**Hit Ratio (taux de bonne direction)** : % des heures où on prédit le **bon sens** (hausse/baisse) du prix. "
  "Important pour le trading.")
H3("Pourquoi on a viré le MAPE")
B("**MAPE** divise par le vrai prix. Quand le prix est proche de 0 (fréquent en élec !), on divise par ~0 → "
  "la valeur **explose** et devient ininterprétable. → on prend **sMAPE** à la place.")

H2("22.  Le test de Diebold-Mariano (DM) — LE point statistique du sujet")
B("**Problème :** B fait MAE = 16,95 et C fait 16,94. C est « mieux »… mais **0,01 c'est peut-être juste du hasard.** "
  "Comment savoir si la différence est **réelle** ou **due à la chance** ?")
B("**Réponse : un test statistique.** Le **test de Diebold-Mariano** compare les erreurs de 2 modèles, heure par "
  "heure, et dit si l'un est **significativement** meilleur.")
H3("Le vocabulaire à maîtriser absolument")
B("**Hypothèse nulle (H0)** : « les 2 modèles sont aussi bons l'un que l'autre » (pas de vraie différence).")
B("**p-value** : la probabilité d'observer une différence aussi grande **si H0 était vraie** (= si en réalité "
  "c'était du hasard). **p petit (< 0,05) → la différence est réelle** (on rejette H0). **p grand → on ne peut pas "
  "conclure** à une différence (on garde H0).")
B("**Notre résultat clé : C vs B → p = 0,572** → **TRÈS au-dessus de 0,05** → on **ne peut pas dire** que la météo "
  "aide. C'est ça « la météo n'est pas significative ».")
B("**Correction HLN (Harvey-Leybourne-Newbold)** : un petit ajustement du test pour le rendre **plus fiable sur "
  "des prévisions** (corrige un biais quand l'échantillon est fini). À citer, ça fait sérieux.")
B("**Statistique DM négative** = le 1er modèle est meilleur. Ex : C vs A = **−53,5** → C écrase le naïf, et "
  "p < 0,001 → ultra significatif.")

H2("23.  L'importance des variables (feature importance / MDI)")
B("**MDI (Mean Decrease in Impurity, baisse moyenne d'impureté)** : mesure, dans la forêt, **combien chaque "
  "variable aide à réduire l'erreur** quand on l'utilise pour couper. Plus une variable est utilisée utilement, "
  "plus son MDI est haut.")
B("**Ce que ça donne (à connaître) :** lag-24h = **20,7 %** (la n°1) ; lags + moyennes glissantes ≈ **70 %** ; "
  "**TTF gaz 8,3 %**, charbon 6 %, ratio nucléaire ~1,9 % ; **TOUTE la météo réunie ≈ 2,2 %.**")
B("→ Ça **confirme** le test DM : la météo pèse très peu, parce que son info est déjà ailleurs.")

# ===================== PARTIE VI =====================
H1("PARTIE VI — LES RÉSULTATS ET LEUR INTERPRÉTATION")

H2("24.  Résultats en régime stable (mai 2024 → avril 2025)")
B("**ML vs naïf :** MAE passe de **33,1 → ~16,9 €/MWh** = **−49 %**. R² de 0,16 → **0,78**. → le ML marche, "
  "confirmé par DM à **p < 0,001**.")
B("**Météo (C vs B) :** MAE bouge de **0,01 €/MWh**. R² de 0,001. → **négligeable.** DM : **p = 0,572 → non "
  "significatif.** ⭐ **C'est LE résultat central de la thèse.**")
B("**RF vs XGBoost :** XGBoost (MAE 19,1) est **moins bon** que RF ici, et significativement (p < 0,001).")

H2("25.  L'hypothèse de redondance d'information (l'explication n°1)")
P("**Pourquoi la météo n'aide pas en régime calme ?** Parce que son signal est **déjà encodé** dans d'autres "
  "variables. On appelle ça l'**information redundancy hypothesis**. 3 canaux :")
B("**1. La prévision de charge ≡ la température.** La load forecast de RTE est **construite à partir de prévisions "
  "de température**. En l'incluant, on conditionne **déjà** sur la météo-demande → la température brute devient "
  "**redondante**.")
B("**2. Le TTF gaz ≡ la météo européenne.** Le prix du gaz monte avec la demande de chauffage de **toute l'Europe**. "
  "Le marché du gaz **agrège déjà** le signal météo.")
B("**3. Domination du nucléaire.** Le ratio de disponibilité nucléaire explique déjà une grosse part de la "
  "variance des prix, indépendamment de la météo.")
callout("La phrase qui tue :",
        "« La météo n'est pas inutile dans l'absolu — elle est **redondante** : l'info utile est déjà captée par "
        "la prévision de charge et le prix du gaz. »")

H2("26.  La validation crise 2022 (LE retournement — notre nouveauté)")
B("**Le risque d'une thèse :** conclure « la météo est inutile » à partir d'**un seul régime** (calme) → ça ne "
  "généralise peut-être pas. Donc on **teste sur un autre régime** : la crise 2022.")
B("**Contexte 2022 :** nucléaire tombé à **~40 %** (corrosion + maintenance, plus bas historique), gaz TTF > "
  "**200 €/MWh** (guerre/coupure russe), prix élec régulièrement > 500–1000 €/MWh.")
B("**Méthode :** on **ré-entraîne sur 2018–2021** et on teste sur **toute 2022** (jamais vue). Split strict.")
B("**Résultat 1 — tout se dégrade :** MAE du modèle C passe de 16,9 → **66,6 €/MWh** (~×4). Les pics de crise "
  "sont largement imprévisibles.")
B("**Résultat 2 — LA MÉTÉO DEVIENT SIGNIFICATIVE :** C vs B → **p < 0,001** (DM = **−13,27**), alors que c'était "
  "p = 0,572 en régime calme. ⭐ **C'est notre trouvaille la plus originale.**")
B("**Résultat 3 :** XGBoost se dégrade **encore plus** que RF (+14,6 %) → confirme que **RF est le plus robuste**.")

H2("27.  Pourquoi la crise casse la redondance (3 mécanismes — à connaître par cœur)")
B("**1. Le TTF se découple de la météo locale.** En 2022 le gaz est piloté par la **géopolitique** (coupure russe), "
  "pas par la température française. → le canal « gaz ≡ météo » **saute**.")
B("**2. Les contraintes nucléaires amplifient la demande.** Avec le nucléaire à ~40 %, **plus de marge** pour "
  "absorber un pic de froid → la température se répercute **directement** sur le prix.")
B("**3. Le rapport signal/bruit monte.** Quand les prix varient de centaines d'€, l'effet direct température→demande "
  "devient **assez gros pour être mesurable** statistiquement.")
callout("Le concept à vendre au jury :",
        "**« regime-dependent » (dépendant du régime).** La valeur de la météo **n'est pas fixe** : "
        "négligeable en régime stable, significative en crise. → Donc inclure la météo doit être une **décision "
        "dynamique**, pas un choix permanent. Recommandation pratique : garder la météo + un **détecteur de régime** "
        "(sur la dispo nucléaire et la corrélation gaz-température) qui l'active en cas de stress.")

# ===================== PARTIE VII =====================
H1("PARTIE VII — LA VALEUR ÉCONOMIQUE (le backtest de trading)")

H2("28.  Pourquoi un backtest ? (stat ≠ argent)")
B("Un bon MAE ne veut pas **automatiquement** dire « ça rapporte ». Il faut vérifier que la précision se traduit "
  "en **bonnes décisions de trading**. → on **simule une stratégie** (backtest) sur le marché réel.")

H2("29.  Le signal de trading (la décision)")
B("Pour chaque heure de demain, règle : **si le modèle prévoit un prix PLUS HAUT qu'aujourd'hui à la même heure "
  "→ on achète (LONG, +1). S'il prévoit plus bas → on vend (SHORT, −1).** Formellement : "
  "signe(prix_prévu_demain − prix_aujourd'hui).")
B("C'est **exécutable** : à 12h aujourd'hui on connaît déjà le prix d'aujourd'hui → pas de triche.")

H2("30.  Le « dead-band » (bande morte) et les coûts")
B("**Dead-band (bande morte) = δ = 2 €/MWh.** Si le mouvement prévu est **tout petit** (< 2 €/MWh), **on ne "
  "trade pas** (position à plat). Pourquoi ? Parce que les frais mangeraient le mini-profit. Filtre de bon sens.")
B("**Coûts de transaction (transaction costs)** : chaque trade paie un frais à la bourse. On teste **3 scénarios** : "
  "**optimiste 0,10**, **central 0,30**, **pessimiste 0,60 €/MWh**. → montre qu'on est rentable même en pessimiste.")

H2("31.  Les métriques de risque (à savoir expliquer)")
B("**Net P&L (Profit & Loss, profit net)** : combien on gagne au total, en **€/MW** sur la période.")
B("**Sharpe Ratio** : rendement **rapporté au risque** (rendement moyen ÷ volatilité). Plus c'est haut, mieux "
  "c'est. On l'**annualise** sur du P&L **journalier** (×√365).")
B("**Maximum Drawdown (MDD, perte max)** : la **pire chute** cumulée du capital (du sommet au creux). "
  "**Petit MDD = stratégie peu risquée** (on ne perd jamais gros longtemps). **Métrique de risque clé.**")
B("**Calmar Ratio** : rendement annualisé ÷ MDD. Mesure le rendement **par unité de risque de perte**. Haut = bien.")
B("**Profit Factor (PF)** : total des gains ÷ total des pertes. > 1 = rentable ; > 2 = bon edge.")
B("**Win Rate** : % d'heures gagnantes.")

H2("32.  Lecture des résultats (scénario central 0,30 €/MWh)")
B("**RF (B et C) ≈ 136 700 €/MW** de profit net, contre **85 500** pour le naïf, et **~0** pour « long-only ».")
B("**Drawdown RF ≈ 420 €/MW** = **5,6× plus petit** que le naïf (2 367). → bien moins risqué.")
B("**B ≈ C** sur TOUTES les métriques → **confirme l'ablation météo** côté économique (la météo n'ajoute rien "
  "au trading en régime stable).")
B("**XGBoost** : P&L un peu plus haut (142 000) MAIS **drawdown 3× pire** → sur base risque, RF gagne.")
B("**Sharpe positif TOUS les 12 mois** → pas un coup de chance d'un seul mois.")
H3("Le « long-only » : pourquoi c'est important")
B("**Long-only** = stratégie bête qui **achète tout le temps**. Elle gagne **~0 €**. → **Preuve que nos profits "
  "viennent du \"skill\" de prévision, pas d'une tendance haussière du marché.** Argument anti-« vous avez juste "
  "eu un marché qui monte ».")

H2("33.  Le piège du Sharpe à 19 (à anticiper absolument)")
B("Un Sharpe de ~19 est **énorme** comparé aux fonds actions (où 2 est déjà très bon). **Ne PAS le vendre comme "
  "ça.**")
B("À dire : c'est du P&L journalier annualisé, **à 1 MW, sans levier, sans taux sans risque déduit**. **Ce n'est "
  "PAS comparable à un fonds actions.** Ce qui compte c'est la **comparaison ENTRE modèles**, et le **drawdown "
  "5,6× plus petit** — pas la valeur absolue du Sharpe. On suppose aussi des **exécutions parfaites** (limite).")

# ===================== PARTIE VIII =====================
H1("PARTIE VIII — CONCLUSIONS, RECOMMANDATIONS, LIMITES")

H2("34.  Les 4 conclusions")
NB("**Le ML bat le naïf** nettement : −49 % MAE, R² 0,78, p < 0,001, robuste sur les 12 mois.", 1)
NB("**La météo est regime-dependent** : non significative en stable (redondance), significative en crise 2022 "
   "(la redondance casse). → **contribution principale.**", 2)
NB("**RF > XGBoost** sur base risque (Calmar 328 vs 107), et RF se dégrade moins en crise.", 3)
NB("**Les prévisions créent de la valeur réelle** : ~137 k€/MW, Sharpe ~19, drawdown minuscule, robuste aux coûts.", 4)

H2("35.  Recommandations pratiques")
B("**Utiliser Random Forest** comme modèle principal (robuste + interprétable).")
B("**Prioriser les fondamentaux en régime calme, garder la capacité météo pour les crises** : implémenter un "
  "**détecteur de régime** qui active/désactive le poids de la météo selon l'état du marché (dispo nucléaire, "
  "découplage gaz-météo).")
B("**Ré-entraîner régulièrement** (mensuel/trimestriel) pour suivre l'évolution du parc et du renouvelable.")
B("**Ajouter des intervalles de prévision** (prediction intervals) pour gérer le risque sur les pics.")

H2("36.  Les limites (montre ton honnêteté intellectuelle — le jury adore)")
B("**Fenêtre de test courte** : 12 mois d'un seul régime stable (on compense avec 2022, mais plus de régimes = mieux).")
B("**Exécution parfaite supposée** dans le backtest : pas de slippage, pas d'impact de marché. À grande taille, "
  "l'edge se réduirait.")
B("**ERA5 = météo parfaite**, pas une vraie prévision (voir encadré Partie III).")
B("**EUA/carbone pas inclus** séparément (assumé, signal dans le spread gaz-charbon).")
B("**XGBoost pas optimisé** finement (réglages par défaut) → l'écart avec RF pourrait se réduire avec du tuning.")

# ===================== PARTIE IX — GLOSSAIRE =====================
H1("PARTIE IX — GLOSSAIRE COMPLET (toutes les abréviations)")
gloss = [
    ("DA", "Day-Ahead — marché « pour le lendemain »."),
    ("EPEX SPOT", "La bourse de l'électricité (France/DE...). Notre marché cible."),
    ("RTE", "Réseau de Transport d'Électricité — le gestionnaire du réseau français (publie la prévision de charge)."),
    ("TSO", "Transmission System Operator — gestionnaire de réseau (RTE en est un)."),
    ("ENTSO-E", "Réseau européen des TSO — notre source de données prix/charge/production."),
    ("ERA5", "Données météo de réanalyse de l'ECMWF (Copernicus)."),
    ("TTF", "Title Transfer Facility — prix de référence du gaz naturel européen (€/MWh)."),
    ("ARA", "Amsterdam-Rotterdam-Anvers — référence prix du charbon (€/t)."),
    ("EUA / ETS", "EU Allowance / Emissions Trading System — quotas carbone (PAS inclus séparément)."),
    ("MW / MWh", "Mégawatt (puissance) / Mégawattheure (énergie). Prix en €/MWh."),
    ("GW", "Gigawatt = 1000 MW (échelle du parc : ~63 GW nucléaire)."),
    ("Load forecast", "Prévision de charge (demande) publiée par RTE la veille — feature clé."),
    ("Merit order", "Ordre de mérite : empilement des centrales du moins au plus cher."),
    ("HDD", "Heating Degree Days — degrés-jours de chauffage (froid sous 17°C)."),
    ("CDD", "Cooling Degree Days — degrés-jours de clim (chaud au-dessus de 22°C)."),
    ("WSI", "Weather Stress Index = HDD / vent (froid + pas de vent)."),
    ("Feature", "Variable explicative donnée au modèle (on en a 35)."),
    ("Lag", "Retard : valeur du prix il y a X heures (lag-24h, lag-168h=1 semaine)."),
    ("Rolling mean", "Moyenne glissante des dernières heures."),
    ("RF", "Random Forest — forêt de 500 arbres moyennés. Notre modèle principal."),
    ("Bagging / Bootstrap", "Entraîner chaque arbre sur un tirage aléatoire avec remise."),
    ("XGBoost", "Modèle de boosting (arbres séquentiels qui se corrigent)."),
    ("Overfitting", "Surapprentissage : coller au passé, mal généraliser."),
    ("Train / Test", "Jeu d'entraînement (passé) / jeu de test (futur, jamais vu)."),
    ("Data leakage", "Fuite de données : le modèle « voit » le futur → à éviter absolument."),
    ("MAE", "Mean Absolute Error — erreur absolue moyenne (€/MWh). Métrique principale."),
    ("RMSE", "Root Mean Square Error — pénalise les grosses erreurs (pics)."),
    ("R²", "Part de variance expliquée (0 = nul, 1 = parfait)."),
    ("sMAPE", "Erreur en % symétrique (évite la division par ~0)."),
    ("Hit Ratio", "% d'heures où on prédit le bon sens (hausse/baisse)."),
    ("DM", "Diebold-Mariano — test qui dit si un modèle est significativement meilleur."),
    ("HLN", "Harvey-Leybourne-Newbold — correction qui fiabilise le test DM."),
    ("H0", "Hypothèse nulle : « les 2 modèles sont aussi bons » (pas de différence)."),
    ("p-value", "Proba que la différence soit du hasard. < 0,05 = différence réelle."),
    ("MDI", "Mean Decrease in Impurity — mesure d'importance des variables dans RF."),
    ("Ablation", "Enlever une chose (la météo) pour mesurer son impact (C vs B)."),
    ("Backtest", "Simuler une stratégie de trading sur des données passées."),
    ("P&L", "Profit & Loss — gain/perte (en €/MW)."),
    ("Dead-band", "Bande morte : on ne trade pas si le mouvement prévu < 2 €/MWh."),
    ("Sharpe", "Rendement / risque (volatilité). Plus haut = mieux."),
    ("MDD / Drawdown", "Maximum Drawdown — pire chute cumulée. Petit = peu risqué."),
    ("Calmar", "Rendement annualisé / MDD."),
    ("Profit Factor", "Total gains / total pertes (>1 rentable)."),
    ("Long-only", "Stratégie qui achète tout le temps (gagne ~0 → preuve de skill)."),
    ("Regime-dependent", "Dépendant du régime de marché — notre concept clé sur la météo."),
]
for ab, de in gloss:
    p = doc.add_paragraph(style="List Bullet"); p.space_after = Pt(1)
    r = p.add_run(ab + " : "); r.bold = True
    p.add_run(de)

# ===================== PARTIE X — A SAVOIR PAR COEUR =====================
H1("PARTIE X — À SAVOIR PAR CŒUR (le minimum vital)")

H2("Les 12 chiffres à connaître")
for t in [
    "Données : 2018–2025, horaire, ~64 000 obs, 35 features, 4 sources.",
    "France : ~63 GW nucléaire (~70 % prod), thermosensibilité ~2,4 GW/°C.",
    "Régime stable : MAE naïf 33,1 → RF ~16,9 €/MWh (−49 %). R² 0,16 → 0,78.",
    "Météo stable (C vs B) : DM = −0,565, p = 0,572 → NON significatif.",
    "Crise 2022 : MAE RF 66,6 (~×4), R² 0,475.",
    "Météo crise (C vs B) : DM = −13,27, p < 0,001 → SIGNIFICATIF (le retournement).",
    "Importance : lag-24h 20,7 % ; lags+rolling ~70 % ; TTF 8,3 % ; météo totale ~2,2 %.",
    "Backtest RF : ~136 700 €/MW, Sharpe ~19,5, MDD ~420.",
    "Naïf backtest : 85 500 €/MW, MDD 2 367 (5,6× pire).",
    "Long-only ≈ 0 €/MW → les profits = skill, pas tendance.",
    "Seuils : HDD 17°C, CDD 22°C, dead-band 2 €/MWh, coûts 0,10/0,30/0,60.",
    "RF : 500 arbres, profondeur 10, min-leaf 5, règle sqrt.",
]:
    B(t)

H2("Les 5 confusions à NE PAS faire")
B("**« La météo est inutile »** → FAUX. Dire : **« redondante en régime stable, significative en crise ».**")
B("**« Notre Sharpe de 19 est exceptionnel »** → PIÈGE. Dire : **non comparable aux actions, regarder le drawdown.**")
B("**« ERA5 est une prévision météo »** → NON, c'est de la **réanalyse** (météo observée) → on teste une météo parfaite.")
B("**« XGBoost est meilleur car plus moderne »** → NON, ici **RF est plus robuste**, surtout en crise.")
B("**« On a triché en utilisant les prix d'aujourd'hui »** → NON, ils sont connus avant 12h (gate closure) → exécutable.")

H2("La structure de l'exposé (qui dit quoi)")
B("**LÉO (≈10 min)** : intro/motivation → question → data → les 4 modèles → résultats stables → finding central "
  "(météo non significative). Passe la main sur « pourquoi ? ».")
B("**LYAM (≈10 min)** : importance des variables → hypothèse de redondance → retournement 2022 → pourquoi la crise "
  "casse la redondance → backtest trading → conclusions.")

callout("Si tu retiens UNE seule chose :",
        "**« La valeur de la météo pour prédire le prix de l'élec français dépend du régime de marché : "
        "elle est redondante en temps normal — car déjà contenue dans la prévision de charge et le prix du gaz — "
        "mais redevient significative en crise, quand le choc gazier et les contraintes nucléaires cassent cette "
        "redondance. »**")

os.makedirs("outputs", exist_ok=True)
doc.save("outputs/Guide_Comprehension_Complet.docx")
print("Saved outputs/Guide_Comprehension_Complet.docx")
