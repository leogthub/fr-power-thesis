# -*- coding: utf-8 -*-
"""Generate Interrogatoire_Blanc.docx from the mock interview content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x00, 0x22, 0x55)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GOLD  = RGBColor(0xC8, 0xA0, 0x32)
GREY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1A, 0x73, 0x4A)
ORANGE = RGBColor(0xB4, 0x5F, 0x1E)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def run(p, text, bold=False, italic=False, color=None, size=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    if size:  r.font.size = Pt(size)
    return r

def h1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(18)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), '002255')
    pBdr.append(bot); pPr.append(pBdr)
    return p

def h2(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(13)
    return p

def body(text, italic=False, color=None, indent=0.4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(indent)
    r = p.add_run(text)
    r.italic = italic
    if color: r.font.color.rgb = color
    r.font.size = Pt(11)
    return p

def q_block(num, question):
    """Question header — bold red number + question text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run(p, f"Q{num} — ", bold=True, color=RED, size=12)
    run(p, question, bold=True, color=NAVY, size=12)
    return p

def answer_label():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    run(p, "Réponse :", bold=True, color=GREEN, size=11)
    return p

def tip_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    run(p, text, bold=True, color=ORANGE, size=11)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot); pPr.append(pBdr)
    return p

def add_recovery_table(items):
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    col_widths = [Cm(2.2), Cm(12.0)]
    for ri, (num, text) in enumerate(items):
        row = table.rows[ri]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = col_widths[0]; c1.width = col_widths[1]
        # number cell — navy bg
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(num)
        r0.bold = True; r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r0.font.size = Pt(12)
        tc0 = c0._tc; tcPr0 = tc0.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'002255')
        tcPr0.append(shd)
        # text cell
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(text)
        r1.italic = True; r1.font.size = Pt(11)
        fill = 'EEF3F9' if ri % 2 == 0 else 'FFFFFF'
        tc1 = c1._tc; tcPr1 = tc1.get_or_add_tcPr()
        shd1 = OxmlElement('w:shd')
        shd1.set(qn('w:val'),'clear'); shd1.set(qn('w:color'),'auto'); shd1.set(qn('w:fill'),fill)
        tcPr1.append(shd1)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════
h1("INTERROGATOIRE BLANC — SOUTENANCE")
p = doc.add_paragraph()
run(p, "The Role of Weather in French Day-Ahead Electricity Price Forecasting", bold=True, color=NAVY, size=12)
p = doc.add_paragraph()
run(p, "Leo Cambreleng & Lyam Oumedjeber  ·  EDHEC MSc DAAI  ·  Juin 2026", color=GREY, size=10)
p = doc.add_paragraph()
run(p, "Format réponse : 1 phrase directe  →  justification  →  chiffre clé  ·  Cible : < 2 min par question",
    italic=True, color=GREY, size=10)
divider()

# ══════════════════════════════════════════════════════════════
# PARTIE A — 20 QUESTIONS
# ══════════════════════════════════════════════════════════════
h1("PARTIE A — 20 Questions avec réponses")

# Q1
q_block(1, "Votre résultat dit que la météo « ne sert pas ». N'est-ce pas décevant ?")
answer_label()
body("Non — c'est une conclusion précise, pas négative. Le vrai résultat a deux parties : redondance en régime stable (p = 0,572), mais significativité en crise 2022 (p < 0,001, DM = −13,27). Une thèse qui dit « parfois oui, parfois non, et voici le mécanisme » est scientifiquement plus riche. L'implication opérationnelle directe : il faut un détecteur de régime, pas une inclusion ou exclusion systématique.")
divider()

# Q2
q_block(2, "Pourquoi Random Forest plutôt qu'un LSTM ou un Transformer ?")
answer_label()
body("Trois raisons. RF capture les non-linéarités à seuil (HDD 17°C) via ses splits. Il est robuste aux outliers (prix > 1 000 EUR/MWh) car il moyenne sur 500 arbres bootstrappés. Et la mesure MDI nous donne l'importance par feature — essentielle pour répondre à notre question de recherche. Un LSTM ne fournit pas cette interprétabilité native.")
divider()

# Q3
q_block(3, "Le test DM avec correction HLN — pourquoi cette correction ?")
answer_label()
body("HLN (Harvey-Leybourne-Newbold, 1997) est un ajustement pour échantillon fini. La statistique DM est asymptotiquement normale, mais sur des horizons courts les queues sont plus épaisses. HLN remplace la loi normale par une loi de Student dont les degrés de liberté dépendent de n. Sur n = 8 640, la différence est faible mais réelle, et la correction rend le test plus conservateur — ce qui renforce la non-significativité en régime stable.")
divider()

# Q4
q_block(4, "ERA5 est une réanalyse, pas une prévision. N'introduisez-vous pas un biais look-ahead ?")
answer_label()
body("Point important. ERA5 donne la météo réelle passée — la « météo parfaite ». En production, on n'aurait que des prévisions NWP avec 10–30% d'erreur. Notre design teste donc la valeur maximale théorique de la météo — une borne supérieure. Si même la météo parfaite est non-significative en régime stable (p = 0,572), une vraie prévision imparfaite l'aurait été encore moins. En 2022, DM = −13,27 laisse suffisamment de marge pour que le résultat tienne malgré l'imprécision NWP.")
divider()

# Q5
q_block(5, "Votre Sharpe de 19,5 est irréaliste pour un vrai fonds. Pourquoi le présentez-vous ?")
answer_label()
body("Ce Sharpe est une mesure relative entre modèles, pas une comparaison avec des fonds d'actions. 1 MW, sans levier, sans déduction du taux sans risque — c'est la convention marché énergie. La métrique pertinente est le Calmar Ratio : RF à 328 vs naïf à 37. Et le drawdown : 422 EUR/MW vs 2 367 pour le naïf. Le Sharpe positif dans tous les 12 mois montre la consistance, pas un niveau absolu.")
divider()

# Q6
q_block(6, "Vous n'avez pas inclus les prix EUA (quotas carbone). N'est-ce pas un facteur déterminant ?")
answer_label()
body("Nous le reconnaissons en section 6.3. L'EUA influence le coût marginal du gaz et du charbon, qui sont eux-mêmes inclus (TTF, ARA coal). Le spread gaz-charbon capture déjà une grande partie du signal carbone. Notre ablation C vs B ne change que les features météo — même avec EUA ajouté, la comparaison météo/sans-météo resterait valide. Ajouter EUA est une extension naturelle mais ne changerait probablement pas notre conclusion principale.")
divider()

# Q7
q_block(7, "XGBoost est généralement meilleur que RF dans les benchmarks. Votre résultat est-il fiable ?")
answer_label()
body("Oui, avec une nuance. L'infériorité d'XGBoost reflète nos hyperparamètres par défaut, pas une infériorité intrinsèque. Le boosting séquentiel sur-pondère les résidus des heures extrêmes (> 500 EUR/MWh en 2022 dans le train), créant de l'overfitting sur ces outliers. RF dilue cet effet en moyennant sur 500 bootstrap. Un tuning Bayésien sur XGBoost pourrait fermer ou inverser le gap. La comparaison B vs C reste propre car même architecture.")
divider()

# Q8
q_block(8, "La « thermosensibilité de 2,4 GW/°C » — d'où vient ce chiffre ?")
answer_label()
body("Ce chiffre est publié par RTE dans le Bilan Électrique 2023. C'est la sensibilité de la demande totale française à 1°C de variation, mesurée empiriquement sur les mois d'hiver. RTE est l'opérateur du réseau, il a accès à toutes les données de consommation. C'est la thermosensibilité la plus élevée d'Europe car environ 35% du chauffage français est électrique, contre ~8% en Allemagne.")
divider()

# Q9
q_block(9, "Votre ablation C vs B n'est valide que si les modèles sont identiques par ailleurs. Comment vous l'assurez-vous ?")
answer_label()
body("Par construction : mêmes hyperparamètres (500 arbres, profondeur 10, min-leaf 5, random_state=42), même train set (jan. 2018 – avr. 2024), même test set (mai 2024 – avr. 2025), même critère d'évaluation. Seule différence : 27 features pour B, 35 pour C (+ 8 variables météo). Le DM compare les erreurs heure par heure sur les mêmes 8 640 observations. C'est un design d'ablation propre — l'équivalent d'une expérience contrôlée.")
divider()

# Q10
q_block(10, "Pourquoi le lag-168h pour le benchmark naïf plutôt que le lag-24h plus courant ?")
answer_label()
body("Le lag-168h est plus exigeant, donc plus honnête. Les prix électriques ont une autocorrélation hebdomadaire forte : le prix à 8h lundi suit mieux le lundi précédent que le dimanche précédent. Le lag-24h est une barre basse. Battre le lag-168h avec −49% de MAE est un résultat plus robuste. C'est la pratique standard en EPF (Weron 2014, Lago 2021).")
divider()

# Q11
q_block(11, "Les prix négatifs représentent 3–8% des observations. Vos modèles les traitent-ils correctement ?")
answer_label()
body("Partiellement. MAE est symétrique et traite les prix négatifs comme toute observation. Nous utilisons sMAPE plutôt que MAPE pour éviter la division par des valeurs proches de zéro. Mais RF cible l'espérance conditionnelle — il ne modélise pas spécifiquement le régime négatif, où la formation des prix est différente (producteurs inflexibles paient pour maintenir l'équilibre). Un modèle dédié à ce régime est mentionné comme extension (section 6.3).")
divider()

# Q12
q_block(12, "Votre dataset commence en 2018. Pourquoi pas 2015 ou 2010 ?")
answer_label()
body("Deux raisons. L'API Copernicus CDS pour ERA5 à résolution horaire sur la France est pratiquement exploitable à partir de 2018. Et la structure du marché français a évolué — le parc solaire et éolien a significativement augmenté avant 2018. Partir plus tôt aurait inclus une distribution des prix non-stationnaire biaisée par une pénétration des renouvelables beaucoup plus faible. Ce choix est cohérent avec la littérature récente (Tschora et al. 2022, Lago et al. 2021).")
divider()

# Q13
q_block(13, "Le lag-24h a 20,7% d'importance mais votre modèle n'est pas simplement « le prix d'hier » ?")
answer_label()
body("Exact. Le lag-24h étant la feature la plus importante ne signifie pas que le modèle est naïf — cela signifie qu'il est bien calibré. La valeur ajoutée du RF vient de la combinaison de toutes les features : 70% lags/rolling means, 8,3% TTF, 6% charbon, etc. L'importance MDI mesure la contribution marginale étant donné les autres features, pas la corrélation univariée. Un modèle « simplement lag-24h » serait notre benchmark — et il est largement battu.")
divider()

# Q14
q_block(14, "En 2022, MAE passe de 16,94 à 66,55. Votre modèle est-il vraiment utile dans ce contexte ?")
answer_label()
body("Oui — la comparaison pertinente est vs le benchmark. En 2022, le naïf atteint 72,74 EUR/MWh. Notre RF à 66,55 représente quand même −8,5%. Le marché est intrinsèquement imprévisible lors d'un choc géopolitique — aucun modèle de régression ne prédit des événements idiosyncratiques. Ce qui compte : nos modèles restent les meilleurs disponibles, le R² passe de 0,160 (naïf) à 0,475 (RF), et en trading le drawdown reste bien contrôlé.")
divider()

# Q15
q_block(15, "La stratégie de trading suppose des exécutions parfaites. Comment les coûts implicites réduisent-ils la performance ?")
answer_label()
body("Nous modélisons les coûts explicites EPEX (0,10 à 0,60 EUR/MWh). Pour 1 MW, le market impact est négligeable sur EPEX — marché très liquide. Pour des positions plus importantes (50–100 MW), il y aurait du cost compression. Mais même à 0,60 EUR/MWh (scénario pessimiste), le Sharpe passe de 19,51 à 19,16 — dégradation < 2%. La robustesse aux coûts est démontrée dans la fourchette réaliste.")
divider()

# Q16
q_block(16, "Avez-vous testé d'autres marchés européens pour valider la généralisation ?")
answer_label()
body("Pas dans cette étude — c'est délibéré et constitue notre principale extension future. L'hypothèse de redondance est spécifique à la structure française : nucléaire dominant (~70%), thermosensibilité électrique élevée, TTF comme proxy météo efficace. En Allemagne, les renouvelables dépassent 60% de la production — le canal météo-offre est beaucoup plus fort. Notre hypothèse prédit que la météo serait plus significative sur EPEX Allemagne qu'en France, même en régime stable.")
divider()

# Q17
q_block(17, "La correction HLN s'applique à des erreurs à horizon fixe. Vos 8 640 observations ne sont pas indépendantes. Le test est-il valide ?")
answer_label()
body("Oui. DM est conçu pour des séries avec autocorrélation. La statistique est basée sur les différentiels de perte d_t = |e_t^C| − |e_t^B|, qui peuvent être autocorrélés. Diebold & Mariano (1995) et Harvey et al. (1997) utilisent un estimateur de variance robuste à l'autocorrélation via un noyau Bartlett (type Newey-West). Le test tient compte de la dépendance sérielle par construction.")
divider()

# Q18
q_block(18, "Que faudrait-il pour reproduire exactement vos résultats ?")
answer_label()
body("Quatre éléments. (1) Données : tout vient de sources publiques — ENTSO-E, Copernicus CDS ERA5, TTF, ARA coal — avec les paramètres d'appel API précisés. (2) Code : disponible dans notre repo GitHub. (3) Graines aléatoires : random_state=42 pour RF et XGBoost. (4) Découpe temporelle : train jan. 2018–avr. 2024 ; test stable mai 2024–avr. 2025 ; test crise train 2018–2021, test 2022.")
divider()

# Q19
q_block(19, "La métrique sMAPE — pourquoi pas simplement le RMSE plus standard ?")
answer_label()
body("Nous reportons le RMSE mais ne l'utilisons pas comme critère principal. RMSE pénalise quadratiquement les grandes erreurs — il est dominé par les heures de prix extrêmes qui représentent < 1% des observations. MAE est plus représentatif de la performance médiane. sMAPE est un indicateur de pourcentage robuste aux valeurs proches de zéro — MAPE pur diverge quand le prix réel → 0 ou négatif, ce qui arrive 3–8% des heures. Tous les cinq métriques sont reportés en Table 4.1.")
divider()

# Q20
q_block(20, "Si vous recommencez, que changeriez-vous en priorité ?")
answer_label()
body("Trois choses. (1) Walk-forward retraining mensuel sur fenêtre expansive — plus proche du production. (2) Optimisation Bayésienne d'XGBoost avec temporal cross-validation — comparaison RF/XGBoost plus honnête. (3) Inclure les prix EUA et tester sur 2023 comme deuxième régime stable. Ces extensions ne remettent pas en cause la conclusion principale — DM = −13,27 est un effet massif qui survivrait à tout raffinement raisonnable.")
divider()

# ══════════════════════════════════════════════════════════════
# PARTIE B — 5 QUESTIONS PIÈGES
# ══════════════════════════════════════════════════════════════
h1("PARTIE B — 5 Questions pièges à gérer à deux", color=RED)
body("Ces questions peuvent aller dans plusieurs directions. Concertez-vous avant — ou répartissez-vous la réponse.", italic=True, color=GREY)
doc.add_paragraph()

h2("B1 — « Vous n'avez testé qu'ERA5. Avez-vous essayé ARPEGE ou des NWP ? »", color=RED)
tip_label("Pourquoi c'est piégeux :")
body("Dire « non, seulement ERA5 » semble non justifié. Dire « oui » sans l'avoir fait est faux.")
tip_label("Comment gérer à deux :")
body("Reconnaître clairement ERA5 uniquement. Argument ERA5 = borne supérieure (Leo, côté technique). Si la météo parfaite est n.s. en régime stable, une NWP imparfaite l'aurait été encore moins (Lyam, implication).")
divider()

h2("B2 — « Comment gérez-vous la non-stationnarité sur 7 ans ? »", color=RED)
tip_label("Pourquoi c'est piégeux :")
body("RF n'a pas de « coefficients » au sens linéaire. La stationnarité pour un RF demande de passer par les importances ou les performances rolling.")
tip_label("Comment gérer à deux :")
body("Lyam : Sharpe mensuel positif tous les 12 mois = stabilité relative. Leo : stationnarité des features (les lags de prix capturent des autocorrélations robustes au-delà du niveau des prix). Mentionner le rolling MAE par mois (Figure 4.4) comme proxy de stabilité.")
divider()

h2("B3 — « Avez-vous envisagé un clustering de régimes pour détecter les crises automatiquement ? »", color=RED)
tip_label("Pourquoi c'est piégeux :")
body("C'est une extension naturelle non réalisée. Ni prétendre l'avoir faite, ni rejeter l'idée.")
tip_label("Comment gérer à deux :")
body("Lyam reformule positivement : c'est exactement notre recommandation finale (ch. 6.1) — un détecteur HMM ou k-means sur la disponibilité nucléaire + volatilité TTF. Leo développe l'opérationnalisation : signaux observables (nuclear_avail_ratio < 0,55, volatilité TTF).")
divider()

h2("B4 — « 0,73 EUR/MWh de gain en crise — est-ce économiquement significatif ? »", color=RED)
tip_label("Pourquoi c'est piégeux :")
body("0,73 EUR/MWh sur une MAE de 67, c'est 1,1% en valeur relative. Pourtant DM = −13,27 avec p < 0,001. Le jury cherche la distinction significativité statistique vs économique.")
tip_label("Comment gérer à deux :")
body("Leo : côté statistique — DM = −13,27 est un effet massif par rapport au bruit de mesure. Lyam : côté économique — en crise à 100–200 EUR/MWh, le gain moyen de 0,73 EUR/MWh masque des moments où l'avantage est bien plus grand. Et le changement de régime (n.s. → ***) est le résultat, pas la magnitude absolue.")
divider()

h2("B5 — « Un modèle dédié météo (SARIMA-GARCH + météo) n'aurait-il pas été plus puissant ? »", color=RED)
tip_label("Pourquoi c'est piégeux :")
body("Arguments valides des deux côtés. Le jury teste si vous avez pesé l'alternative.")
tip_label("Comment gérer à deux :")
body("Leo défend le choix : l'ablation B vs C est valide précisément parce que l'architecture est identique — elle isole la contribution des features. Comparer C (RF) à un SARIMA-météo comparerait architecture ET features simultanément, rendant l'ablation impossible. Lyam : Lago et al. (2021) montrent que RF et modèles économétriques ont des performances comparables sur les marchés européens.")
divider()

# ══════════════════════════════════════════════════════════════
# PARTIE C — 10 PHRASES DE RÉCUPÉRATION
# ══════════════════════════════════════════════════════════════
h1("PARTIE C — 10 Phrases de récupération")
body("Pour les blancs, les surprises ou les attaques rhétoriques. Prononce calmement, gagne du temps, reviens à tes chiffres.", italic=True, color=GREY)
doc.add_paragraph()

add_recovery_table([
    ("1", "Je ne sais pas → \"That's a great question — let me think for a second so I give you a precise answer rather than an approximation.\""),
    ("2", "Je confonds un chiffre → \"I want to give you the exact figure — the value reported in our paper is [pause] — rather than risk citing an approximation on a central statistic.\""),
    ("3", "Je ne comprends pas la question → \"To make sure I answer what you're asking: is your question mainly about [A] or about [B]?\""),
    ("4", "Attaque sur une limite → \"You're absolutely right to flag that — we document it ourselves in section 6.3. The question is whether it invalidates the main result, and we believe it doesn't, for the following reason...\""),
    ("5", "Passer la main → \"Leo/Lyam worked on this specific aspect — I'll hand over for the technical detail.\""),
    ("6", "Question très large → \"That touches several aspects of our methodology. On [aspect 1]: ...; on [aspect 2]: ...\""),
    ("7", "Erreur dans la présentation → \"Thank you for flagging that — the correct value is [Y], as reported in Table [Z].\""),
    ("8", "Hors périmètre → \"That goes beyond our study's scope — we deliberately limited ourselves to [France / RF+XGBoost / 2018-2025]. It's precisely why we propose [extension] as future work.\""),
    ("9", "Conclure n'importe quelle réponse → \"In summary: our central result — regime-dependence of weather value — stands. [Key figure]. [Mechanism in one sentence].\""),
    ("10", "Silence gênant → \"That's exactly the question that pushed us to run the 2022 validation — let me explain how we answered it empirically.\""),
])

# ══════════════════════════════════════════════════════════════
# CHIFFRES CLÉS (aide-mémoire final)
# ══════════════════════════════════════════════════════════════
h1("CHIFFRES CLÉS — à connaître par cœur")

table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
facts = [
    ("MAE stable C vs B", "16,94 vs 16,95 EUR/MWh — diff. = 0,01"),
    ("DM stable (C vs B)", "stat = −0,565  ·  p = 0,572  →  n.s."),
    ("DM crise 2022 (C vs B)", "stat = −13,27  ·  p < 0,001  →  ***"),
    ("ML vs naïf (stable)", "−49% MAE  ·  33,09 → 16,94  ·  p < 0,001"),
    ("Feature importance météo", "~2,2% de l'importance totale"),
    ("P&L RF (central 0,30 EUR/MWh)", "~137 000 EUR/MW/an"),
    ("MaxDD RF vs naïf", "422 EUR vs 2 367 EUR — ratio 5,6×"),
    ("Calmar RF vs naïf", "328 vs 37"),
    ("Sharpe RF (tous scénarios)", "19,51 → 19,46 → 19,16 — dégradation < 2%"),
]
for ri, (label, val) in enumerate(facts):
    row = table.rows[ri]
    c0, c1 = row.cells[0], row.cells[1]
    c0.width = Cm(5.5); c1.width = Cm(8.7)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    r0.bold = True; r0.font.color.rgb = NAVY; r0.font.size = Pt(11)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(val)
    r1.font.size = Pt(11)
    fill = 'EEF3F9' if ri % 2 == 0 else 'FFFFFF'
    for cell, f in [(c0, 'EEF3F9'), (c1, fill)]:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), f)
        tcPr.append(shd)

doc.add_paragraph()
p = doc.add_paragraph()
run(p, "Bon courage — la thèse est solide, les chiffres parlent d'eux-mêmes.", italic=True, color=GREY, size=11)

# ── Save ──────────────────────────────────────────────────────
OUT = r"C:\Users\Public\fr-power-thesis\outputs\Interrogatoire_Blanc.docx"
doc.save(OUT)
print("Saved:", OUT)
