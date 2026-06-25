# -*- coding: utf-8 -*-
"""Build the defense support Word document: timed script, Q&A, 1-page cheat sheet."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EDHEC = RGBColor(0, 62, 126)
ACCENT = RGBColor(46, 95, 163)
GREY = RGBColor(110, 110, 110)
GREEN = RGBColor(34, 110, 60)
RED = RGBColor(170, 45, 35)

doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)


def shade(par, hex_color):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hex_color)
    pPr.append(sh)


def h1(text, color=EDHEC):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = color
    p.space_before = Pt(10); p.space_after = Pt(4)
    return p


def h2(text, color=ACCENT):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = color
    p.space_before = Pt(8); p.space_after = Pt(2)
    return p


def body(runs, size=10.5, after=4, before=0):
    p = doc.add_paragraph(); p.space_after = Pt(after); p.space_before = Pt(before)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, ov in runs:
        r = p.add_run(text)
        r.font.size = Pt(ov.get("size", size))
        r.bold = ov.get("bold", False); r.italic = ov.get("italic", False)
        if "color" in ov: r.font.color.rgb = ov["color"]
    return p


def bullet(runs, size=10.5):
    p = doc.add_paragraph(style="List Bullet"); p.space_after = Pt(2)
    if isinstance(runs, str): runs = [(runs, {})]
    for text, ov in runs:
        r = p.add_run(text); r.font.size = Pt(ov.get("size", size))
        r.bold = ov.get("bold", False); r.italic = ov.get("italic", False)
        if "color" in ov: r.font.color.rgb = ov["color"]
    return p


def hr():
    p = doc.add_paragraph(); p.space_before = Pt(4); p.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '2E5FA3')
    pb.append(bottom); pPr.append(pb)


# ===== TITLE BLOCK =====
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ORAL DEFENSE — SUPPORT PACK"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = EDHEC
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
r = p.add_run("The Role of Weather in French Day-Ahead Electricity Price Forecasting")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(6)
r = p.add_run("Leo Cambreleng & Lyam Oumedjeber  ·  Supervisor: Prof. Milos Vulanovic  ·  EDHEC MSc Data Analysis & AI")
r.font.size = Pt(9.5); r.font.color.rgb = GREY
hr()
body([("Defense format: ", {"bold": True}),
      ("~60 min = 20-min presentation + 20–40 min Q&A (online via Teams). Target ~1.3 min/slide, 14 slides. "
       "Proposed split below — adjust freely.", {})], size=10)

# ===== PART 1: TIMED SCRIPT =====
h1("Part 1 — Timed Speaker Script")
script = [
    ("1", "Title", "BOTH", "0:00–0:30", [
        "Good morning, thank you for being here. We are Leo Cambreleng and Lyam Oumedjeber, MSc Data Analysis & AI.",
        "Our thesis asks a deceptively simple question: does weather actually help forecast French day-ahead electricity prices? The answer turned out to be: it depends on the regime.",
    ]),
    ("2", "Context & Motivation", "LÉO", "0:30–2:30", [
        "Electricity is unique: it can't be stored, so supply and demand must clear every single hour — that produces extreme volatility and spikes.",
        "France is special: ~63 GW of nuclear, ~70% of output, and a heating stock so electric that demand moves ~2.4 GW per °C in winter — the highest thermosensitivity in Europe.",
        "This matters economically: on a 100 MW book, improving accuracy by just 1 EUR/MWh is worth ~876,000 EUR a year.",
        "So the natural question: temperature and wind clearly drive demand and renewables — but is that signal already priced into the data we feed the model?",
    ]),
    ("3", "Research Question", "LÉO", "2:30–4:00", [
        "Our central question is on screen. We break it into four sub-questions: (1) can ML beat a naïve benchmark; (2) do weather features add significant value; (3) does that depend on the regime; (4) does it translate into trading profit.",
        "The third one — regime dependence — is where our main contribution lies.",
    ]),
    ("4", "Data", "LÉO", "4:00–5:30", [
        "We built a reproducible hourly dataset, 2018 to 2025, ~64,000 observations, from four public sources: ENTSO-E for prices/load/generation/flows, ERA5 reanalysis for weather, plus TTF gas and ARA coal.",
        "From these we engineer 35 features across five families.",
        "The price series on the right shows three regimes: calm pre-2021, the 2022 crisis above 1,000 EUR/MWh, and normalisation after. That crisis becomes our natural experiment later.",
    ]),
    ("5", "Four Models", "LÉO", "5:30–7:00", [
        "We use a nested four-model design. Model A is a naïve lag-168h benchmark — last week's same hour. B is a Random Forest on 27 features with NO weather. C is the same RF plus full weather, 35 features — our main model. D is XGBoost on the same 35 features.",
        "The key is the C-versus-B comparison: it isolates exactly the marginal value of weather. We judge it with the Diebold-Mariano test, HLN-corrected.",
    ]),
    ("6", "Results — Stable Regime", "LÉO", "7:00–9:00", [
        "On the 2024–25 test set: both Random Forests cut MAE by about 49% versus naïve — from 33 down to ~17 EUR/MWh — and lift R² from 0.16 to 0.78. Machine learning clearly works.",
        "But look at C versus B: MAE moves by one eurocent. Adding weather does essentially nothing here.",
        "And XGBoost actually underperforms the Random Forest with these settings.",
    ]),
    ("7", "Central Finding", "LÉO → hands to Lyam", "9:00–10:30", [
        "Formally: the DM test for weather, C vs B, gives p = 0.572 — we cannot reject equal accuracy. Weather is not significant.",
        "Meanwhile both RFs beat naïve at p < 0.001, and XGBoost is significantly worse than RF.",
        "So ML adds robust value — but weather, here, does not. The obvious question is WHY, and whether that's always true. Lyam will take it from here.",
    ]),
    ("8", "Feature Importance", "LYAM", "10:30–12:00", [
        "Thanks Leo. The importances tell the story. Price lags dominate — the 24-hour lag alone is 20.7%, and lags plus rolling means are about 70% of all importance.",
        "Fuel prices matter too: TTF gas 8.3%, coal 6%. All weather features combined are only ~2.2%.",
        "That's the mechanism behind the non-significance — the weather signal seems already captured elsewhere.",
    ]),
    ("9", "Redundancy Hypothesis", "LYAM", "12:00–13:30", [
        "We call this the information redundancy hypothesis. Three channels: first, the RTE load forecast is BUILT on temperature — so including it already conditions on weather-driven demand.",
        "Second, TTF gas prices rise with pan-European heating demand — the commodity market aggregates the weather signal for us.",
        "Third, nuclear availability explains a lot of price variance on its own. So raw weather becomes statistically redundant.",
    ]),
    ("10", "2022 Crisis Reversal", "LYAM", "13:30–15:30", [
        "Now the key test. A finding from one calm regime may not generalise — so we retrain on 2018–2021 and test on the full 2022 crisis, strict temporal split.",
        "The result reverses completely. Weather, C vs B, goes from p = 0.572 to p < 0.001, DM −13.27 — highly significant.",
        "Everything degrades — MAE quadruples to ~67 — but weather now genuinely helps, and Random Forest is far more robust than XGBoost.",
    ]),
    ("11", "Why It Breaks", "LYAM", "15:30–17:00", [
        "Why? The same three channels break. One: the gas crisis was geopolitical, so TTF decoupled from French temperature. Two: with nuclear at ~40%, there's no headroom — cold demand feeds straight into price. Three: when prices swing by hundreds of euros, the direct temperature effect is finally large enough to measure.",
        "The takeaway: weather's value is regime-dependent — a dynamic modelling choice, not a permanent one.",
    ]),
    ("12", "Trading Backtest", "LYAM", "17:00–18:30", [
        "Does accuracy mean money? We run an executable day-ahead directional strategy with realistic costs.",
        "The Random Forests make ~137,000 EUR per MW versus 85,000 for naïve, and a long-only benchmark makes essentially zero — so this is genuine skill, not market drift.",
        "Crucially, RF drawdown is 5.6× smaller than naïve, with a positive Sharpe in every single month.",
    ]),
    ("13", "Conclusions", "LYAM", "18:30–19:45", [
        "Four takeaways: ML beats naïve decisively; weather is regime-dependent — our core contribution; Random Forest beats XGBoost on risk-adjusted terms; and the forecasts create real, robust economic value.",
        "Our recommendation: use weather features with a regime-detection switch — prioritise fundamentals when calm, activate weather in crises.",
    ]),
    ("14", "Thank You / Q&A", "BOTH", "19:45–20:00", [
        "That concludes our presentation. Thank you — we'd be glad to take your questions.",
    ]),
]
for num, title, who, timing, lines in script:
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(2)
    shade(p, "EAF0F7")
    r = p.add_run(f"  Slide {num} · {title}"); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = EDHEC
    r2 = p.add_run(f"      [{who}]  {timing}"); r2.bold = True; r2.font.size = Pt(9.5); r2.font.color.rgb = ACCENT
    for ln in lines:
        bullet(ln, size=10)

# ===== PART 2: Q&A =====
doc.add_page_break()
h1("Part 2 — Anticipated Questions & Answers")
body([("Milos will probe method, data, and interpretation. Lead with a one-line answer, then justify. "
       "Below: the question, then your prepared answer.", {"italic": True, "color": GREY})], size=10, after=6)

qa = [
    ("Isn't a p = 0.572 'weather doesn't matter' just an underpowered test or bad features?",
     "No — three convergent pieces of evidence point the same way: the DM test (p=0.572), the feature importances (all weather = 2.2%), and the trading backtest (B and C give near-identical P&L). And the test clearly HAS power — it detects every other effect at p<0.001 on the same 8,640 hours. The null isn't 'no signal in weather', it's 'no signal beyond what load and TTF already encode'."),
    ("Why Random Forest rather than a neural net / LSTM / LEAR?",
     "RF fits the problem: it captures threshold non-linearities (e.g. the 17°C heating kink), is robust to the extreme price outliers common in power markets, and gives interpretable MDI importances — which we needed precisely because our research question is about feature contribution. LSTMs/transformers are a stated extension; with ~64k rows and our interpretability goal, RF is the right baseline."),
    ("XGBoost usually beats Random Forest. Why doesn't it here?",
     "With default-ish hyperparameters XGBoost's sequential boosting over-weights extreme training observations, so it's more fragile to outliers and regime shifts — its 2022 degradation (+14.6%) confirms this. We're explicit that proper Bayesian hyperparameter tuning could close or reverse the gap; we didn't tune it to keep the comparison fair and reproducible."),
    ("Your load forecast already contains temperature — isn't including it circular / leaking weather?",
     "It's not leakage — the RTE load forecast is genuinely available before the day-ahead auction closes, so it's legitimate information. The point is the opposite of circular: it's exactly WHY raw weather looks redundant. The load forecast is a better-engineered weather feature than our raw ERA5 means, because RTE's proprietary model already maps temperature to demand."),
    ("ERA5 is reanalysis — it uses observed weather, not a forecast. Isn't that look-ahead bias?",
     "Fair and important. ERA5 is the actual realised weather, so we're testing the value of PERFECT weather information — an upper bound. If even perfect weather is redundant in the stable regime, a real forecast can only be weaker, which strengthens our conclusion. For the crisis result it's a caveat we state: real NWP forecast error would shrink, but not erase, the gain."),
    ("A Sharpe of ~19 is absurd for a real strategy. Is the backtest realistic?",
     "We flag this explicitly. That number is daily P&L annualised at 1 MW with no leverage and no risk-free deduction — it is NOT comparable to an equity fund. The relevant comparison is between models on the same scale. We also assume perfect fills and ignore market impact; at size, the edge would compress. The robust claim is the RANKING and the 5.6× smaller drawdown, not the absolute Sharpe."),
    ("Only 12 months of test data — could the stable-regime result be luck?",
     "We address this two ways: positive Sharpe in all 12 months and a cost-robust result rule out a single lucky month; and the 2022 out-of-sample validation tests an entirely different regime. The honest limitation is that we have one stable window — more regimes would strengthen generalisation, which we list as future work."),
    ("You dropped EU ETS carbon prices — doesn't that bias the fundamentals model?",
     "We acknowledge it as a limitation. The carbon signal is largely embedded in the TTF–coal spread we do include, so the marginal-cost ordering of gas vs coal is still captured. Adding EUA explicitly is a clean, low-cost extension; we don't expect it to change the weather conclusion, since it would reinforce, not replace, the fundamental channel."),
    ("How do you actually USE 'regime-dependence' in practice?",
     "Operationally: keep weather features in the model but add a regime detector on observable signals — nuclear availability ratio and the gas–temperature correlation. When those flag stress (low nuclear, decoupled gas), up-weight weather; in calm regimes, lean on fundamentals. It turns a binary 'include or not' into a state-dependent decision."),
    ("What's genuinely novel here versus the existing EPF literature?",
     "Three things: a complete reproducible French dataset; a formal weather-ablation with DM testing specifically for France; and — the core novelty — showing the weather effect is regime-dependent, using 2022 as a natural experiment and explaining the mechanism via redundancy that breaks under supply shocks. Most EPF work treats weather value as fixed; we show it isn't."),
    ("Negative and near-zero prices — how did you handle them in the metrics?",
     "We use MAE as the primary metric (robust to outliers) and deliberately drop MAPE, which blows up near zero. We report sMAPE instead, which bounds the denominator. RMSE is reported too, to show spike sensitivity — and the RMSE/MAE gap is exactly why we recommend prediction intervals for risk management."),
    ("If weather barely helps, why include it at all?",
     "Two reasons. First, in the stable regime it's costless to keep — it doesn't hurt accuracy. Second, and decisively, the 2022 result shows it provides significant, measurable value precisely when you most need accuracy — during a crisis. Dropping it would leave you exposed exactly when prices are most violent."),
]
for i, (q, a) in enumerate(qa, 1):
    body([(f"Q{i}.  ", {"bold": True, "color": EDHEC, "size": 10.5}), (q, {"bold": True, "size": 10.5})], after=2, before=6)
    body([("A.  ", {"bold": True, "color": GREEN, "size": 10}), (a, {"size": 10})], after=2)

# ===== PART 3: ONE-PAGE CHEAT SHEET =====
doc.add_page_break()
h1("Part 3 — One-Page Cheat Sheet")
body([("Keep this in front of you. All numbers you might be asked for.", {"italic": True, "color": GREY})], size=10, after=6)

h2("The one-sentence thesis")
body([("The marginal value of weather in French day-ahead price forecasting is regime-dependent: "
       "redundant in stable conditions (load forecast + TTF already encode it), but significant in the 2022 crisis "
       "when the gas shock and nuclear constraints break that redundancy.", {"bold": True})], size=10.5, after=6)

h2("Headline numbers — stable regime (May 2024–Apr 2025, n=8,640)")
for t in [
    "Naïve A: MAE 33.09 · R² 0.160", "RF no-wx B: MAE 16.95 · R² 0.775",
    "RF wx C: MAE 16.94 · R² 0.776 (main)", "XGBoost D: MAE 19.14 · R² 0.659",
    "ML vs naïve: −49% MAE · DM p<0.001", "Weather C vs B: DM −0.565 · p=0.572 (n.s.)",
]:
    bullet(t, size=10)

h2("Headline numbers — 2022 crisis (n=8,760)")
for t in [
    "RF wx C: MAE 66.55 · R² 0.475 (≈4× worse)",
    "Weather C vs B: DM −13.27 · p<0.001 (***)  ← the reversal",
    "XGBoost degrades +14.6% vs RF (crisis fragility)",
]:
    bullet(t, size=10)

h2("Backtest — central cost 0.30 EUR/MWh, 1 MW")
for t in [
    "RF: ~136,700 EUR/MW · Sharpe ~19.5 · MaxDD ~420",
    "Naïve: 85,482 EUR/MW · MaxDD 2,367 (5.6× larger)",
    "Long-only: ≈ 0 → returns are skill, not trend",
    "Positive Sharpe in all 12 months; robust to costs (drag <2%)",
]:
    bullet(t, size=10)

h2("Feature importance (RF wx)")
for t in [
    "lag-24h 20.7% · lags+rolling ≈ 70%",
    "TTF gas 8.3% · coal 6.0% · nuclear ratio ~1.9%",
    "ALL weather combined ≈ 2.2%",
]:
    bullet(t, size=10)

h2("Setup facts")
for t in [
    "Data: 2018–2025 hourly, ~64k obs, 4 public sources (ENTSO-E, ERA5, TTF, coal); 35 features",
    "Models: A naïve lag-168h · B RF 27f · C RF 35f · D XGBoost 35f",
    "Test: DM with Harvey-Leybourne-Newbold correction, MAE-based loss",
    "RF: 500 trees, depth 10, min-leaf 5, sqrt features",
]:
    bullet(t, size=10)

h2("If you blank — the 4 conclusions")
for t in [
    "1. ML beats naïve (−49% MAE, p<0.001)",
    "2. Weather is regime-dependent (n.s. stable → *** crisis)",
    "3. RF > XGBoost on risk-adjusted basis (Calmar 328 vs 107)",
    "4. Forecasts create real value (~137k EUR/MW, tiny drawdown)",
]:
    bullet([(t, {"bold": True})], size=10)

os.makedirs("outputs", exist_ok=True)
doc.save("outputs/Defense_Support.docx")
print("Saved outputs/Defense_Support.docx")
